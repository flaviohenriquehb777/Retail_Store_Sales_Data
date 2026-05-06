from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> Path:
    docx_path = Path("reports") / "Relatorio_Executivo_Recuperacao_de_Margem_Acentuado.docx"
    if not docx_path.exists():
        raise FileNotFoundError(str(docx_path))

    doc = Document(docx_path)

    doc.add_page_break()
    doc.add_heading("Respostas Dirtas Para as Perguntas", level=1)

    doc.add_heading("1. Por que a margem de lucro está caindo mesmo com o aumento do faturamento?", level=2)
    for b in [
        "A empresa está crescendo faturamento com deterioração do “preço efetivo” (preço – desconto + frete) e com aumento do peso de vendas em recortes de baixa rentabilidade.",
        "Há incidência elevada de pedidos com prejuízo (lucro < 0), o que puxa o resultado consolidado para baixo mesmo quando a receita cresce.",
        "A queda está concentrada em pontos específicos do portfólio e em interações com a operação logística (principalmente Transporte Rodoviário e, em alguns focos, Aéreo Normal).",
        "Em 2026, a categoria Mobiliário tornou-se estruturalmente deficitária, com destaque para Mesas e Estantes, o que explica grande parte da compressão de margem no consolidado.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2. Como podemos reverter a queda da margem de lucro?", level=2)
    doc.add_paragraph(
        "A reversão exige combinar governança comercial (guardrails), correção do preço efetivo e ajustes operacionais. "
        "Somente otimizar mix não é suficiente, pois parte relevante do faturamento está ocorrendo em condições que geram prejuízo."
    )
    doc.add_paragraph("Ações obrigatórias e imediatas:", style="List Bullet")
    for b in [
        "Implementar guardrails: bloquear pedidos com lucro < 0 (ou exigir aprovação) e exigir margem mínima por pedido (>=10% em Aéreo; >=12% em Rodoviário).",
        "Regras de desconto por subcategoria (teto): eliminar descontos nos focos que geram prejuízo recorrente e reduzir descontos em recortes que tiveram forte erosão de margem.",
        "Regras de frete: Rodoviário somente com repasse de frete e/ou margem mínima; restringir Aéreo Rápido/Normal em focos com margem negativa quando houver subsídio de frete.",
    ]:
        doc.add_paragraph(b, style="List Bullet 2")
    doc.add_paragraph("Ações prescritivas por foco (onde atuar):", style="List Bullet")
    for b in [
        "Mobiliário > Mesas e Estantes: zerar desconto, impedir frete subsidiado em Rodoviário, e reprecificar para retirar a subcategoria do prejuízo.",
        "Material de Escritório > Armazenamento e Organização: eliminar faixa 5–10% de desconto e restringir modalidades de envio subsidiadas nos recortes deficitários.",
        "Tecnologia: reduzir concessões na faixa 5–10% de desconto (especialmente em Máquinas de Escritório e Periféricos), mantendo competitividade apenas onde a margem suporta.",
        "Nordeste e Pequenas Empresas (quando Rodoviário): endurecer regras de desconto e repasse de frete, pois são recortes com deterioração relevante.",
    ]:
        doc.add_paragraph(b, style="List Bullet 2")

    doc.add_heading("3. Quais são os produtos que estão diminuindo a margem da empresa?", level=2)
    doc.add_paragraph("Produtos/subcategorias com maior contribuição para queda de margem (principalmente em 2026):")
    for b in [
        "Mobiliário > Mesas (margem negativa em 2026; maior peso em faturamento e maior drenagem de lucro).",
        "Mobiliário > Estantes (margem fortemente negativa em 2026).",
        "Material de Escritório > Armazenamento e Organização (margem negativa em 2026, com concentração de prejuízo em combinações de envio/condições comerciais).",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph("Combinações críticas (produto x envio) que agravaram a queda:")
    for b in [
        "Mesas + Transporte Rodoviário.",
        "Estantes + Transporte Rodoviário.",
        "Armazenamento e Organização + Aéreo Normal (e também sob condições de desconto inadequadas).",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph("Exemplos de itens (SKUs) com prejuízo relevante em 2026 (amostra):")
    for b in [
        "Okidata Pacemark 4410N Wide Format Dot Matrix Printer.",
        "Epson DFX5000+ Dot Matrix Printer.",
        "Riverside Palais Royal Lawyers Bookcase (Royale Cherry Finish).",
        "Epson DFX-8500 Dot Matrix Printer.",
        "Chromcraft Bull-Nose Wood Rectangular Conference Tables (48\" x 96\").",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4. Quais produtos focar para aumentar a lucratividade da empresa?", level=2)
    doc.add_paragraph(
        "O foco deve ser duplo: (i) parar imediatamente a venda com prejuízo nos focos críticos e (ii) acelerar crescimento nos focos com maior lucro absoluto e/ou boa margem, "
        "sem voltar a degradar preço efetivo via desconto/frete."
    )
    doc.add_paragraph("Produtos/subcategorias recomendados para foco de crescimento com rentabilidade:", style="List Bullet")
    for b in [
        "Material de Escritório > Capas e Acessórios (alto lucro e alta margem).",
        "Tecnologia > Telefones e Comunicação (alto lucro e boa margem).",
        "Tecnologia > Máquinas de Escritório (alto lucro, porém exige governança de desconto para evitar erosão).",
        "Mobiliário > Cadeiras (bom lucro e margem positiva).",
        "Tecnologia > Periféricos (lucro e margem positivos; manter desconto sob controle).",
        "Tecnologia > Copiadoras e fax (lucro e margem positivos; atenção à combinação com Aéreo Normal em condições agressivas).",
        "Material de Escritório > Eletrodomésticos (margem positiva e bom lucro).",
    ]:
        doc.add_paragraph(b, style="List Bullet 2")
    doc.add_paragraph("Como executar (regras de execução para crescer com margem):", style="List Bullet")
    for b in [
        "Aplicar guardrails de margem por pedido e bloquear prejuízo antes de tentar acelerar volume.",
        "Crescer por cross-sell e bundles para elevar ticket e diluir custo de envio (especialmente quando houver Rodoviário).",
        "Usar campanhas de preço apenas em subcategorias com margem comprovadamente suficiente e sempre com limite de desconto por subcategoria.",
    ]:
        doc.add_paragraph(b, style="List Bullet 2")

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    print(main())

