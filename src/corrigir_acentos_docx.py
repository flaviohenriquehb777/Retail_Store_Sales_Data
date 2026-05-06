from __future__ import annotations

import re
from pathlib import Path

from docx import Document

TRANSLATE = str.maketrans(
    {
        "¾": "ó",
        "þ": "ç",
        "Ò": "ã",
        "ß": "á",
        "Ú": "é",
        "Ý": "í",
        "§": "õ",
    }
)

RX_ISOLATED_O_ACUTE = re.compile(r"(?:(?<=\\s)|^)Ó(?:(?=\\s)|$)")


def fix_text(s: str) -> str:
    s = "" if s is None else str(s)
    s = s.translate(TRANSLATE)
    s = RX_ISOLATED_O_ACUTE.sub("à", s)
    return s

def fix_doc(doc: Document) -> None:
    for p in doc.paragraphs:
        for r in p.runs:
            r.text = fix_text(r.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = fix_text(r.text)


def main() -> Path:
    import sys

    docx_path = (
        Path(sys.argv[1])
        if len(sys.argv) > 1
        else (Path("reports") / "Relatorio_Executivo_Recuperacao_de_Margem_Acentuado.docx")
    )
    if not docx_path.exists():
        raise FileNotFoundError(str(docx_path))

    doc = Document(docx_path)
    fix_doc(doc)
    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    print(main())
