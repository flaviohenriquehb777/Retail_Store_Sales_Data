from __future__ import annotations

from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
import unicodedata
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def fmt_brl(x: float) -> str:
    v = float(x or 0.0)
    s = f"{v:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {s}"


def fmt_int(x: float) -> str:
    v = int(round(float(x or 0.0)))
    s = f"{v:,}".replace(",", ".")
    return s


def fmt_pct(x: float, digits: int = 2) -> str:
    if x is None or (isinstance(x, float) and not np.isfinite(x)):
        return "-"
    return f"{x*100:.{digits}f}%"


def fix_text(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    repl = {
        "Escrit¾rio": "Escritório",
        "Mobilißrio": "Mobiliário",
        "OrganizaþÒo": "Organização",
        "Acess¾rios": "Acessórios",
        "EletrodomÚsticos": "Eletrodomésticos",
        "Elßsticos": "Elásticos",
        "rÚguas": "réguas",
        "aparadores": "aparadores",
        "ComunicaþÒo": "Comunicação",
        "Mßquinas": "Máquinas",
        "PerifÚricos": "Periféricos",
        "Armazenamento e OrganizaþÒo": "Armazenamento e Organização",
        "Capas e Acess¾rios": "Capas e Acessórios",
        "Tesouras, rÚguas e aparadores": "Tesouras, réguas e aparadores",
        "Telefones e ComunicaþÒo": "Telefones e Comunicação",
        "Mobilißrio de Escrit¾rio": "Mobiliário de Escritório",
        "Material de Escrit¾rio": "Material de Escritório",
        "AÚreo": "Aéreo",
        "NÒo": "Não",
        "MÚdia": "Média",
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    return s


def norm_ascii(s: str) -> str:
    s = "" if s is None else str(s)
    s = unicodedata.normalize("NFKD", s)
    s = s.encode("ascii", "ignore").decode("ascii")
    return " ".join(s.split())


def disc_bin(x: float) -> str:
    x = float(x or 0.0)
    if x <= 0:
        return "0%"
    if x <= 0.05:
        return "0-5%"
    if x <= 0.10:
        return "5-10%"
    if x <= 0.20:
        return "10-20%"
    if x <= 0.30:
        return "20-30%"
    return "30%+"


def agg(df: pd.DataFrame, by: list[str]) -> pd.DataFrame:
    g = (
        df.groupby(by)
        .agg(
            faturamento=("Faturamento", "sum"),
            lucro=("Lucro", "sum"),
            custo_envio=("Custo de Envio", "sum"),
            pedidos=("Order ID", "nunique"),
            linhas=("Row ID", "count"),
        )
        .reset_index()
    )
    g["margem"] = np.where(g["faturamento"] != 0, g["lucro"] / g["faturamento"], np.nan)
    g["custo_envio_pct_fat"] = np.where(g["faturamento"] != 0, g["custo_envio"] / g["faturamento"], np.nan)
    return g


def delta_2025_2026(df: pd.DataFrame, by: list[str]) -> pd.DataFrame:
    a = agg(df[df["Ano"] == 2025], by).rename(
        columns={
            "faturamento": "fat_2025",
            "lucro": "luc_2025",
            "margem": "marg_2025",
            "pedidos": "ped_2025",
        }
    )
    b = agg(df[df["Ano"] == 2026], by).rename(
        columns={
            "faturamento": "fat_2026",
            "lucro": "luc_2026",
            "margem": "marg_2026",
            "pedidos": "ped_2026",
        }
    )
    out = a.merge(b, on=by, how="outer").fillna(0)
    out["d_fat"] = out["fat_2026"] - out["fat_2025"]
    out["d_luc"] = out["luc_2026"] - out["luc_2025"]
    out["d_marg_pp"] = (out["marg_2026"] - out["marg_2025"]) * 100
    fat_2026_total = float(df[df["Ano"] == 2026]["Faturamento"].sum() or 1.0)
    out["fat_share_2026"] = out["fat_2026"] / fat_2026_total
    return out


def needed_uplift(margem_atual: float, margem_alvo: float) -> float:
    if margem_alvo >= 1:
        return np.nan
    return (margem_alvo - margem_atual) / (1 - margem_alvo)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = val


def main() -> Path:
    base_path = Path("Case 1 - Case Retail Store Sales Data.xlsx")
    df = pd.read_excel(base_path, sheet_name="Sales_retail_store")

    rename_map: dict[str, str] = {}
    for c in df.columns:
        k = norm_ascii(c).lower()
        if k == "preco unitario":
            rename_map[c] = "Preço Unitário"
        if k == "regiao":
            rename_map[c] = "Região"
    if rename_map:
        df = df.rename(columns=rename_map)

    for c in [
        "Prioridade",
        "Forma de Envio",
        "Estado",
        "Região",
        "Segmento do Cliente",
        "Categoria do Produto",
        "Sub-Categoria do Produto",
    ]:
        if c in df.columns:
            df[c] = df[c].astype(str).map(fix_text)

    df["Data da Venda"] = pd.to_datetime(df["Data da Venda"])
    df["Ano"] = df["Data da Venda"].dt.year

    for c in ["Faturamento", "Lucro", "Desconto", "Custo de Envio"]:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)

    df["Faixa de Desconto"] = df["Desconto"].map(disc_bin)

    fat_total = float(df["Faturamento"].sum())
    luc_total = float(df["Lucro"].sum())
    margem_total = luc_total / fat_total if fat_total else np.nan
    linhas_total = int(df.shape[0])
    pedidos_total = int(df["Order ID"].nunique())
    d0 = df["Data da Venda"].min().date()
    d1 = df["Data da Venda"].max().date()

    anos = agg(df, ["Ano"]).sort_values("Ano")

    order = df.groupby("Order ID", as_index=False).agg(fat=("Faturamento", "sum"), lucro=("Lucro", "sum"))
    neg_orders = order[order["lucro"] < 0]
    neg_share_orders = len(neg_orders) / max(1, len(order))
    neg_fat_share = float(neg_orders["fat"].sum()) / float(order["fat"].sum() or 1.0)

    cat_2526 = delta_2025_2026(df, ["Categoria do Produto"]).sort_values("d_luc")
    ship_2526 = delta_2025_2026(df, ["Forma de Envio"]).sort_values("d_luc")
    reg_2526 = delta_2025_2026(df, ["Região"]).sort_values("d_luc")
    seg_2526 = delta_2025_2026(df, ["Segmento do Cliente"]).sort_values("d_luc")
    sub_2526 = delta_2025_2026(df, ["Categoria do Produto", "Sub-Categoria do Produto"]).sort_values("d_luc")

    foco = {
        "Mesas": ("Mobiliário", "Mesas"),
        "Estantes": ("Mobiliário", "Estantes"),
        "Armazenamento e Organização": ("Material de Escritório", "Armazenamento e Organização"),
        "Máquinas de Escritório": ("Tecnologia", "Máquinas de Escritório"),
        "Periféricos": ("Tecnologia", "Periféricos"),
    }

    margem_alvo = float(anos.loc[anos["Ano"] == 2025, "margem"].iloc[0]) if (anos["Ano"] == 2025).any() else 0.111

    foco_stats = []
    for nome, (cat, sub) in foco.items():
        d = df[(df["Ano"] == 2026) & (df["Categoria do Produto"] == cat) & (df["Sub-Categoria do Produto"] == sub)]
        fat = float(d["Faturamento"].sum())
        luc = float(d["Lucro"].sum())
        marg = luc / fat if fat else np.nan
        uplift = needed_uplift(marg, margem_alvo) if np.isfinite(marg) else np.nan
        foco_stats.append(
            {
                "categoria": cat,
                "subcategoria": sub,
                "faturamento_2026": fat,
                "lucro_2026": luc,
                "margem_2026": marg,
                "uplift_preco": uplift,
            }
        )

    comb_2026 = agg(df[df["Ano"] == 2026], ["Categoria do Produto", "Sub-Categoria do Produto", "Forma de Envio"])
    fat_2026_total = float(df[df["Ano"] == 2026]["Faturamento"].sum() or 1.0)
    comb_2026["share_2026"] = comb_2026["faturamento"] / fat_2026_total
    comb_crit = comb_2026[comb_2026["share_2026"] >= 0.01].sort_values(["margem", "faturamento"], ascending=[True, False]).head(20)

    tech_2026 = df[df["Ano"] == 2026]
    tech_disc = agg(tech_2026[tech_2026["Categoria do Produto"] == "Tecnologia"], ["Faixa de Desconto"]).sort_values("Faixa de Desconto")

    tech_5_10 = tech_2026[(tech_2026["Categoria do Produto"] == "Tecnologia") & (tech_2026["Faixa de Desconto"] == "5-10%")]
    tech_5_10_sub = agg(tech_5_10, ["Sub-Categoria do Produto"]).sort_values("margem")

    lanes = df[df["Ano"] == 2026].groupby(["Order ID", "Região", "Forma de Envio"], as_index=False).agg(
        fat=("Faturamento", "sum"), lucro=("Lucro", "sum"), ship=("Custo de Envio", "sum")
    )
    lanes["marg"] = np.where(lanes["fat"] != 0, lanes["lucro"] / lanes["fat"], np.nan)
    lanes2 = (
        lanes.groupby(["Região", "Forma de Envio"], as_index=False)
        .agg(pedidos=("Order ID", "nunique"), faturamento=("fat", "sum"), lucro=("lucro", "sum"), custo_envio=("ship", "sum"))
        .assign(margem=lambda d: np.where(d["faturamento"] != 0, d["lucro"] / d["faturamento"], np.nan))
        .sort_values("margem")
    )

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph("Relatório Executivo - Diagnóstico e Plano de Recuperação de Margem")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')}\nBase: Case 1 - Retail Store Sales Data")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    doc.add_heading("1. Sumário executivo", level=1)
    p = doc.add_paragraph()
    p.add_run("Problema central: ").bold = True
    p.add_run("o faturamento cresce, mas o lucro não acompanha devido à deterioração de margem e presença significativa de pedidos com prejuízo.")

    bullets = [
        f"Período analisado: {d0.strftime('%d/%m/%Y')} a {d1.strftime('%d/%m/%Y')}.",
        f"Tamanho da base: {fmt_int(linhas_total)} linhas e {fmt_int(pedidos_total)} pedidos.",
        f"Faturamento total: {fmt_brl(fat_total)}; lucro total: {fmt_brl(luc_total)}; margem total: {fmt_pct(margem_total)}.",
        f"Pedidos com prejuízo (lucro < 0): {fmt_pct(neg_share_orders)} dos pedidos; representam {fmt_pct(neg_fat_share)} do faturamento.",
        "Principal alavanca de recuperação: bloquear/prevenir vendas com margem negativa e corrigir preço efetivo (preço, desconto e frete) em pontos específicos do portfólio e operação.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2. Contexto e objetivo", level=1)
    doc.add_paragraph(
        "O objetivo deste relatório é apresentar, de forma acessível aos diretores e sócios, o que os dados indicam sobre o desempenho de vendas e lucratividade, "
        "as principais causas da deterioração de margem e um plano prescritivo (executável) para recuperar a rentabilidade."
    )

    doc.add_heading("3. Base de dados e metodologia", level=1)
    doc.add_paragraph(
        "A análise foi conduzida em três camadas complementares: (i) descritiva, para quantificar o desempenho e sua evolução; "
        "(ii) diagnóstica, para identificar onde e por que a margem se deteriora; e (iii) prescritiva, para definir decisões e regras operacionais que elevem a margem com rapidez."
    )
    add_table(
        doc,
        ["Escopo", "Descrição"],
        [
            ["Unidade de análise", "Linhas de pedido e pedidos (Order ID)."],
            ["Indicadores", "Faturamento, lucro, margem (lucro/faturamento), custo de envio, desconto."],
            ["Cortes analisados", "Ano, mês, categoria, subcategoria, região, segmento, forma de envio, faixa de desconto."],
        ],
    )

    doc.add_heading("4. Análise descritiva - o que aconteceu", level=1)
    doc.add_paragraph("A seguir, os principais resultados agregados por ano:")
    rows_ano = []
    for _, r in anos.iterrows():
        rows_ano.append([str(int(r["Ano"])), fmt_brl(r["faturamento"]), fmt_brl(r["lucro"]), fmt_pct(r["margem"])])
    add_table(doc, ["Ano", "Faturamento", "Lucro", "Margem"], rows_ano)

    doc.add_paragraph(
        "O resultado descritivo confirma o fenômeno reportado pela gestão: crescimento de faturamento não se traduz em crescimento proporcional de lucro, "
        "por conta de deterioração de margem e concentração de vendas em condições de baixa rentabilidade."
    )

    doc.add_heading("5. Análise diagnóstica - por que aconteceu", level=1)
    doc.add_paragraph("Os principais vetores de deterioração identificados foram:")
    for b in [
        "Alta incidência de pedidos com prejuízo, indicando concessões comerciais e/ou custos (especialmente frete) incompatíveis com o preço efetivo.",
        "Deterioração concentrada em pontos específicos do portfólio (categoria/subcategoria) e na operação logística (forma de envio).",
        "Interações críticas entre produto e forma de envio, além de recortes relevantes por região e segmento.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph("Categorias com pior deterioração de lucro entre 2025 e 2026 (visão executiva):")
    cat_focus = cat_2526.sort_values("d_luc").head(3)
    add_table(
        doc,
        ["Categoria", "Fatur. 2025", "Marg. 2025", "Fatur. 2026", "Marg. 2026", "Δ lucro (2026-2025)"],
        [
            [
                fix_text(r["Categoria do Produto"]),
                fmt_brl(r["fat_2025"]),
                fmt_pct(r["marg_2025"]),
                fmt_brl(r["fat_2026"]),
                fmt_pct(r["marg_2026"]),
                fmt_brl(r["d_luc"]),
            ]
            for _, r in cat_focus.iterrows()
        ],
    )

    doc.add_paragraph("Combinações críticas (categoria, subcategoria e envio) com participação relevante em 2026 (≥1% do faturamento de 2026):")
    add_table(
        doc,
        ["Categoria", "Subcategoria", "Envio", "Faturamento 2026", "Lucro 2026", "Margem 2026", "Share 2026"],
        [
            [
                fix_text(r["Categoria do Produto"]),
                fix_text(r["Sub-Categoria do Produto"]),
                fix_text(r["Forma de Envio"]),
                fmt_brl(r["faturamento"]),
                fmt_brl(r["lucro"]),
                fmt_pct(r["margem"]),
                fmt_pct(r["share_2026"], 2),
            ]
            for _, r in comb_crit.iterrows()
        ],
    )

    doc.add_paragraph("Recortes com pior margem em 2026 por região e forma de envio (efeito operacional):")
    lanes_focus = lanes2.head(6)
    add_table(
        doc,
        ["Região", "Envio", "Pedidos", "Faturamento", "Lucro", "Margem"],
        [
            [
                fix_text(r["Região"]),
                fix_text(r["Forma de Envio"]),
                fmt_int(r["pedidos"]),
                fmt_brl(r["faturamento"]),
                fmt_brl(r["lucro"]),
                fmt_pct(r["margem"]),
            ]
            for _, r in lanes_focus.iterrows()
        ],
    )

    doc.add_heading("6. Análise prescritiva - o que fazer (com exatidão)", level=1)

    doc.add_paragraph(
        "O plano prescritivo abaixo foi desenhado para ser operacional: define regras de aprovação comercial, limites de desconto, condições de frete e ajustes de preço "
        "com foco em estancar perdas e recuperar margem de forma rápida."
    )

    doc.add_heading("6.1 Regras obrigatórias (guardrails)", level=2)
    for b in [
        "Bloquear pedidos com lucro do pedido < 0 (ou exigir aprovação gerencial).",
        "Definir margem mínima por pedido: ≥10% (Aéreo Normal/Rápido) e ≥12% (Transporte Rodoviário).",
        "Sem exceções automáticas: exceções somente com justificativa formal e registro de aprovação.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("6.2 Ações por portfólio (categoria e subcategoria)", level=2)

    doc.add_paragraph("Focos prioritários para correção de preço efetivo e política comercial em 2026:")
    add_table(
        doc,
        ["Categoria", "Subcategoria", "Faturamento 2026", "Lucro 2026", "Margem 2026", "Ajuste de preço efetivo (estim.)"],
        [
            [
                s["categoria"],
                s["subcategoria"],
                fmt_brl(s["faturamento_2026"]),
                fmt_brl(s["lucro_2026"]),
                fmt_pct(s["margem_2026"]),
                "-" if not np.isfinite(s["uplift_preco"]) else f"+{s['uplift_preco']*100:.1f}%",
            ]
            for s in foco_stats
        ],
    )

    doc.add_paragraph("Regras executáveis por foco:")
    matriz = [
        [
            "Mobiliário > Mesas",
            "Desconto máximo 0%. Transporte Rodoviário somente com frete repassado e margem mínima ≥12%.",
            "Reprecificar: aumentar preço efetivo até recuperar a margem-alvo (estimativa acima).",
        ],
        [
            "Mobiliário > Estantes",
            "Desconto máximo 0%. Transporte Rodoviário somente com frete repassado e margem mínima ≥12%.",
            "Reprecificar: aumentar preço efetivo até recuperar a margem-alvo (estimativa acima).",
        ],
        [
            "Material de Escritório > Armazenamento e Organização",
            "Eliminar desconto 5–10%. Aéreo Rápido somente com frete repassado e margem mínima ≥12%.",
            "Reprecificar: aumentar preço efetivo até recuperar a margem-alvo (estimativa acima).",
        ],
        [
            "Tecnologia > Máquinas de Escritório",
            "Eliminar desconto 5–10% (máximo 0–5%) e exigir guardrail de margem por pedido.",
            "Reprecificar seletivamente ou reduzir desconto equivalente, preservando competitividade em itens âncora.",
        ],
        [
            "Tecnologia > Periféricos",
            "Capar desconto em 0–5% e exigir guardrail de margem por pedido.",
            "Ajustar preço/condições comerciais para evitar erosão de margem em volume relevante.",
        ],
    ]
    add_table(doc, ["Onde", "Regra comercial e logística", "Ajuste de preço"], matriz)

    doc.add_heading("6.3 Ações por desconto e envio (política operacional)", level=2)
    doc.add_paragraph("Tecnologia — margem por faixa de desconto (2026):")
    add_table(
        doc,
        ["Faixa de desconto", "Faturamento", "Lucro", "Margem"],
        [[r["Faixa de Desconto"], fmt_brl(r["faturamento"]), fmt_brl(r["lucro"]), fmt_pct(r["margem"])] for _, r in tech_disc.iterrows()],
    )
    doc.add_paragraph("Tecnologia — faixa 5–10%: subcategorias com pior margem (2026):")
    add_table(
        doc,
        ["Subcategoria", "Faturamento", "Lucro", "Margem"],
        [
            [r["Sub-Categoria do Produto"], fmt_brl(r["faturamento"]), fmt_brl(r["lucro"]), fmt_pct(r["margem"])]
            for _, r in tech_5_10_sub.head(4).iterrows()
        ],
    )

    doc.add_heading("6.4 Ações por região e segmento (onde aplicar regras diferenciadas)", level=2)
    doc.add_paragraph("Regiões e segmentos com maior deterioração de lucro entre 2025 e 2026 (prioridade):")
    reg_focus = reg_2526.sort_values("d_luc").head(3)
    seg_focus = seg_2526.sort_values("d_luc").head(3)
    add_table(
        doc,
        ["Recorte", "Fatur. 2025", "Marg. 2025", "Fatur. 2026", "Marg. 2026", "Δ lucro"],
        [
            [f"Região: {r['Região']}", fmt_brl(r["fat_2025"]), fmt_pct(r["marg_2025"]), fmt_brl(r["fat_2026"]), fmt_pct(r["marg_2026"]), fmt_brl(r["d_luc"])]
            for _, r in reg_focus.iterrows()
        ]
        + [
            [
                f"Segmento: {r['Segmento do Cliente']}",
                fmt_brl(r["fat_2025"]),
                fmt_pct(r["marg_2025"]),
                fmt_brl(r["fat_2026"]),
                fmt_pct(r["marg_2026"]),
                fmt_brl(r["d_luc"]),
            ]
            for _, r in seg_focus.iterrows()
        ],
    )
    for b in [
        "Nordeste: proibir frete grátis em Transporte Rodoviário; exigir margem mínima ≥12% ou repasse integral de frete.",
        "Pequenas Empresas: quando Transporte Rodoviário, limitar desconto a 0–5% e exigir margem mínima ≥12% (caso contrário, repassar frete ou bloquear).",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("7. Plano de implantação (30-60-90 dias)", level=1)
    add_table(
        doc,
        ["Janela", "Entregas obrigatórias"],
        [
            [
                "0–30 dias",
                "Implantar guardrails; bloquear pedidos com prejuízo; limitar descontos nos focos; repassar frete em rodoviário nas regiões/segmentos críticos.",
            ],
            [
                "31–60 dias",
                "Reprecificação seletiva (Mesas, Estantes, Armazenamento e Organização); revisão de políticas comerciais por modal; padronização de aprovação de exceções.",
            ],
            [
                "61–90 dias",
                "Otimização de mix (empurrar itens de margem maior); bundles para diluir frete; rotinas de monitoramento e revisão mensal de tabela/preço efetivo.",
            ],
        ],
    )

    doc.add_heading("8. Métricas de sucesso e governança", level=1)
    for b in [
        "Meta 1: reduzir fortemente o percentual de pedidos com lucro < 0 (nível atual: elevado).",
        "Meta 2: recuperar margem total para pelo menos o patamar de 2025 e estabilizar acima dele.",
        "Monitorar semanalmente: margem por (Categoria, Subcategoria, Região, Segmento, Envio, Faixa de desconto) e volume/faturamento em cada recorte.",
        "Governança: toda exceção a desconto/frete deve ser registrada e aprovada, com motivo e prazo de validade.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    out_path = Path("Relatorio_Executivo_Recuperacao_de_Margem.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    p = main()
    print(str(p))

